import warnings
from docx import Document as DocxDocument
from app.ingestion.parsers.base import Page


def parse_docx(path: str) -> list[Page]:
    doc = DocxDocument(path)
    # TODO: demo 范围不提取表格内容；后续接入表格解析（unstructured/marker）
    if doc.tables:
        warnings.warn("docx tables are not extracted (demo scope)", stacklevel=2)
    text = "\n".join(p.text for p in doc.paragraphs)
    return [Page(page_number=1, text=text)]
